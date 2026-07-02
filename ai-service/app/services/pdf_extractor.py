"""
Stage 1 — Document Processing

Extract clean text from PDF/DOCX files (received as bytes from FastAPI UploadFile).

Pipeline:
  1. PyMuPDF smart layout extraction (1-col vs 2-col detection)
  2. Quality scoring (0–100 heuristic)
  3. If PDF quality < 60 → OCR fallback (Tesseract)
  4. DOCX → python-docx paragraph extraction
"""

from __future__ import annotations

import io
import re

import fitz                         # PyMuPDF
from PIL import Image
import pytesseract
from docx import Document           # python-docx


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Dispatch based on file extension.
    Returns clean text (whitespace normalized, control chars removed).
    """
    ext = filename.lower().rsplit(".", 1)[-1]

    if ext == "pdf":
        text = extract_text_smart_layout(file_bytes)
        quality = evaluate_extracted_text_quality(text)
        if quality["score"] < 60:
            # Fallback: rasterize + OCR
            text = ocr_pdf_with_tesseract(file_bytes)
    elif ext in ("docx", "doc"):
        text = extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: .{ext} (only .pdf / .docx allowed)")

    return clean_text(text)


# ---------------------------------------------------------------------------
# PDF — smart layout extraction (port from POC)
# ---------------------------------------------------------------------------

def extract_text_smart_layout(file_bytes: bytes) -> str:
    """
    PyMuPDF block-level extraction with 1-column vs 2-column detection.
    For 2-col layouts, reads header → right column → left column to preserve
    reading order in CVs that put the timeline on the left and content on right.
    """
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    all_pages: list[str] = []

    for page in doc:
        page_width  = page.rect.width
        page_height = page.rect.height
        raw_blocks  = page.get_text("blocks")

        blocks = []
        for block in raw_blocks:
            x0, y0, x1, y1, text, *_ = block
            text = text.strip()
            if not text or len(text) < 2:
                continue
            blocks.append({
                "x0": x0, "y0": y0, "x1": x1, "y1": y1,
                "center_x": (x0 + x1) / 2,
                "text": text,
            })

        if not blocks:
            continue

        blocks.sort(key=lambda b: (b["y0"], b["x0"]))
        left_count  = sum(1 for b in blocks if b["center_x"] < page_width * 0.45)
        right_count = sum(1 for b in blocks if b["center_x"] > page_width * 0.55)
        is_two_col  = left_count >= 2 and right_count >= 2

        parts: list[str] = []
        if not is_two_col:
            parts.extend(b["text"] for b in blocks)
        else:
            header_limit = page_height * 0.16
            header_b = sorted(
                [b for b in blocks if b["y0"] < header_limit],
                key=lambda b: (b["y0"], b["x0"]),
            )
            body_b  = [b for b in blocks if b["y0"] >= header_limit]
            right_b = sorted(
                [b for b in body_b if b["center_x"] >= page_width * 0.45],
                key=lambda b: (b["y0"], b["x0"]),
            )
            left_b  = sorted(
                [b for b in body_b if b["center_x"] < page_width * 0.45],
                key=lambda b: (b["y0"], b["x0"]),
            )
            parts.extend(b["text"] for b in header_b)
            parts.extend(b["text"] for b in right_b)
            parts.extend(b["text"] for b in left_b)

        all_pages.append("\n".join(parts))

    doc.close()
    return "\n\n".join(all_pages)


# ---------------------------------------------------------------------------
# Quality heuristic
# ---------------------------------------------------------------------------

def evaluate_extracted_text_quality(text: str) -> dict:
    """
    Heuristic 0–100 score. Low score = likely scanned/image PDF → need OCR.

    Signals:
      - Length (very short text → poor extraction)
      - Word count (raw chars vs. tokens)
      - Garbage chars ratio (replacement chars, weird symbols)
      - Average word length (sane words are 4–8 chars)
    """
    score = 100
    reasons: list[str] = []

    if len(text) < 100:
        score -= 60
        reasons.append("text too short")

    words = re.findall(r"\b\w+\b", text)
    if len(words) < 30:
        score -= 30
        reasons.append("too few words")

    if text:
        garbage = sum(1 for c in text if c in "�")
        if garbage / max(len(text), 1) > 0.02:
            score -= 20
            reasons.append("high garbage char ratio")

    if words:
        avg_len = sum(len(w) for w in words) / len(words)
        if avg_len < 2 or avg_len > 15:
            score -= 15
            reasons.append(f"abnormal avg word length {avg_len:.1f}")

    return {"score": max(0, score), "reasons": reasons, "word_count": len(words)}


# ---------------------------------------------------------------------------
# OCR fallback
# ---------------------------------------------------------------------------

def ocr_pdf_with_tesseract(file_bytes: bytes, dpi: int = 200, lang: str = "eng+vie") -> str:
    """
    Rasterize each page at `dpi`, run Tesseract OCR.
    Requires Tesseract binary installed (Docker: tesseract-ocr + tesseract-ocr-vie).
    """
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages_text: list[str] = []

    for page in doc:
        # Render page to PNG bytes
        zoom = dpi / 72                  # 72 DPI is PyMuPDF default
        matrix = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=matrix, alpha=False)
        img_bytes = pix.tobytes("png")
        img = Image.open(io.BytesIO(img_bytes))

        page_text = pytesseract.image_to_string(img, lang=lang)
        pages_text.append(page_text)

    doc.close()
    return "\n\n".join(pages_text)


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Read all paragraphs + table cells from a .docx file."""
    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)

    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Text cleaning
# ---------------------------------------------------------------------------

def clean_text(text: str) -> str:
    """Normalize whitespace, strip control characters."""
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
